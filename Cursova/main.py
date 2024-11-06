import comtypes.client
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog
from datetime import datetime

# Функція для конвертації .docx в .pdf за допомогою Word
def convert_docx_to_pdf(docx_path, pdf_path):
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    try:
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 - формат PDF
    finally:
        doc.Close()
        word.Quit()

# Функція для заміни тексту у кожному абзаці з урахуванням форматування
def replace_text_in_paragraphs(paragraphs, data):
    for paragraph in paragraphs:
        # Об'єднуємо весь текст абзацу
        full_text = "".join(run.text for run in paragraph.runs)
        
        # Замінюємо всі ключі в об'єднаному тексті
        for key, value in data.items():
            placeholder = f"{{{key}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)
        
        # Очищаємо поточні `runs`
        for run in paragraph.runs:
            run.clear()
        
        # Вставляємо оновлений текст, додаючи `run`, якщо його не існує
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)

# Функція для заміни тексту у всіх таблицях з урахуванням роздільників
def replace_text_in_tables(tables, data):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, data)

# Функція для формування імені файлу у форматі "Прізвище_Дата народження"
def generate_filename(data):
    # Отримуємо прізвище і дату народження
    surname = data.get("FullName", "").split()[0]
    birth_date = data.get("PassportWhen", "")
    
    # Переводимо дату народження у формат "dd.mm.yyyy"
    try:
        birth_date_obj = datetime.strptime(birth_date, "%d.%m.%Y")
        birth_date_str = birth_date_obj.strftime("%d-%m-%Y")
    except ValueError:
        birth_date_str = "Unknown_Date"
    
    # Формуємо ім'я файлу
    filename = f"{surname}_{birth_date_str}"
    return filename

# Функція для заповнення шаблону даними, включаючи таблиці
def fill_document(data, template_path, result_path):
    # Відкриваємо шаблон Word-документа
    doc = Document(template_path)

    # Заповнюємо текст у головному документі
    replace_text_in_paragraphs(doc.paragraphs, data)

    # Заповнюємо текст у таблицях
    replace_text_in_tables(doc.tables, data)

    # Формуємо ім'я файлу
    filename = generate_filename(data)

    # Збереження документа як .docx у вибрану папку
    doc.save(os.path.join(result_path, f'{filename}.docx'))
    print(f"Word document successfully saved as {filename}.docx")

    # Конвертуємо в PDF після збереження
    convert_docx_to_pdf(os.path.join(result_path, f'{filename}.docx'), os.path.join(result_path, f'{filename}.pdf'))
    print(f"Document successfully converted to {filename}.pdf")

# Дані для заповнення
data = {
    "FullName": "Панчишин Володимир Тарасович",
    "PassportSeries": "",
    "PassportNumber": "123456",
    "PassportWho": "5555",
    "PassportWhen": "15.10.2024",
    "Address": "Місто Львів Вулиця соснова 26/33",
    "RNOKPP": "1234567890",
    "PhoneNumber": "+380123456789",
    "FullName2": "Панчишин Тарас Володимирович",
    "PassportSeries2": "ВК",
    "PassportNumber2": "123456",
    "PassportWho2": "Shevchenka St., 10, Kyiv",
    "PassportWhen2": "12.12.2000",
    "RNOKPP2": "1234567890",
    "Address2": "Вулиця Караджича",
    "PhoneNumber2": "+380123456789"
}

# Вибір папки для збереження результатів
def select_folder():
    root = tk.Tk()
    root.withdraw()  # Приховуємо основне вікно
    folder_path = filedialog.askdirectory(title="Виберіть папку для збереження файлів")
    return folder_path

# Шлях до шаблону
template_path = r'C:\Users\Viva\Desktop\Repositor\Cursova\template.docx'

# Вибір папки для збереження результату
selected_folder = select_folder()

# Перевірка, чи вибрана папка
if selected_folder:
    # Викликаємо функцію для заповнення документа і конвертації в PDF
    fill_document(data, template_path, selected_folder)
else:
    print("Не вибрано папку для збереження.")
